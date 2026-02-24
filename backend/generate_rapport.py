"""
Génération du rapport de conception professionnel au format Word (.docx)
pour le projet Secure Votes — Système de collecte et d'agrégation homomorphe de votes.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ═══════════════════════════════════════════════════
# Styles
# ═══════════════════════════════════════════════════

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x0B, 0x12, 0x24)
    if level == 1:
        heading_style.font.size = Pt(22)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    else:
        heading_style.font.size = Pt(13)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════

def add_colored_heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark blue background
        shading = hdr_cells[i]._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '0B1224',
            qn('w:color'): 'auto',
            qn('w:val'): 'clear'
        })
        shading.append(bg)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def add_code_block(code, language=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    # Background shading
    shading = p._element.get_or_add_pPr()
    bg = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F1F5F9',
        qn('w:color'): 'auto',
        qn('w:val'): 'clear'
    })
    shading.append(bg)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" — {text}")
    else:
        p.add_run(text)
    return p

def add_note_box(text, title="Note"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"⚠️ {title} : ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xB4, 0x5D, 0x09)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.italic = True
    # Background
    shading = p._element.get_or_add_pPr()
    bg = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'FEF3C7',
        qn('w:color'): 'auto',
        qn('w:val'): 'clear'
    })
    shading.append(bg)

# ═══════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Rapport de Conception et Développement")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0B, 0x12, 0x24)

doc.add_paragraph()

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("Système Sécurisé de Collecte et\nd'Agrégation Homomorphe de Votes")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)

doc.add_paragraph()
doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run("Projet CNAM — Sécurité des Données\nFévrier 2026")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc_p.add_run(
    "Chiffrement homomorphe Paillier · AES-256-GCM · PostgreSQL\n"
    "FastAPI · Progressive Web App · Docker Compose · CI/CD"
)
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ═══════════════════════════════════════════════════

doc.add_heading("Table des matières", level=1)

toc_items = [
    "1. Introduction",
    "2. Architecture globale de la solution",
    "3. Modèle UML",
    "4. Choix technologiques",
    "5. Conception cryptographique",
    "6. Sécurité",
    "7. Application mobile (PWA)",
    "8. Backend — API REST",
    "9. Base de données — Séparation en silos",
    "10. Audit et journalisation immuable",
    "11. CI/CD",
    "12. IAM — Gestion des accès",
    "13. Tests et scripts d'audit",
    "14. Conformité et bilan",
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════

doc.add_heading("1. Introduction", level=1)

doc.add_heading("1.1 Contexte", level=2)
doc.add_paragraph(
    "Dans le cadre d'opérations de consultation citoyenne ou de sondages sur le terrain, "
    "des agents se déplacent physiquement pour interroger directement les participants. "
    "Afin de garantir la fiabilité, l'intégrité et la confidentialité des données collectées, "
    "un dispositif numérique sécurisé est nécessaire, capable d'identifier les acteurs, "
    "de contrôler les accès et d'assurer un stockage chiffré des informations sensibles."
)

doc.add_heading("1.2 Objectifs du projet", level=2)
doc.add_paragraph(
    "Le système vise à concevoir une solution mobile + backend sécurisé permettant :"
)
objectives = [
    "La collecte de réponses binaires (Oui/Non) à des questions pré-enregistrées",
    "L'identification fiable des participants et agents via QR codes",
    "Le chiffrement homomorphe des votes (Paillier) pour une agrégation sans déchiffrement",
    "Le chiffrement des photos (AES-256-GCM) côté client",
    "Le stockage dans des silos séparés et chiffrés",
    "La journalisation immuable par chaîne de hachage (hash-chain)",
    "La transmission sécurisée via HTTPS",
    "Un pipeline CI/CD avec linting, tests et scans de sécurité",
]
for obj in objectives:
    add_bullet(obj)

doc.add_heading("1.3 Périmètre technique", level=2)
add_table(
    ["Composant", "Technologie"],
    [
        ["Backend", "Python 3.11 / FastAPI"],
        ["Base de données", "PostgreSQL 15"],
        ["Stockage objet", "MinIO (S3-compatible)"],
        ["Chiffrement votes", "Paillier (homomorphe additif)"],
        ["Chiffrement photos", "AES-256-GCM (Web Crypto API)"],
        ["Application mobile", "Progressive Web App (HTML/JS)"],
        ["Orchestration", "Docker Compose"],
        ["CI/CD", "GitHub Actions"],
        ["Tests", "pytest + httpx"],
    ],
    col_widths=[5, 10]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 2. ARCHITECTURE
# ═══════════════════════════════════════════════════

doc.add_heading("2. Architecture globale de la solution", level=1)

doc.add_heading("2.1 Vue d'ensemble", level=2)
doc.add_paragraph(
    "L'architecture suit un modèle microservices containerisé avec Docker Compose. "
    "Quatre services distincts communiquent via des réseaux logiques séparés :"
)

add_table(
    ["Service", "Rôle", "Port exposé"],
    [
        ["frontend (Nginx)", "Application agent terrain PWA + reverse proxy", "3000"],
        ["api (FastAPI)", "API REST sécurisée, validation, crypto, audit", "8080"],
        ["db (PostgreSQL)", "Stockage en 4 silos séparés", "55432 (dev)"],
        ["minio", "Stockage objet chiffré (photos)", "9000/9001"],
    ],
    col_widths=[4, 8, 3]
)

doc.add_heading("2.2 Schéma d'architecture", level=2)

doc.add_paragraph(
    "L'architecture est organisée en couches successives, du client jusqu'aux données :"
)

architecture_desc = """
┌──────────────────────────────────────────────────────────┐
│                    CLIENT / MOBILE                        │
│  ┌─────────────┐  ┌──────────┐  ┌───────────────────┐   │
│  │  QR Scanner  │  │  Caméra  │  │ Crypto JS         │   │
│  │  (html5-qr)  │  │  Photo   │  │ Paillier + AES-GCM│   │
│  └─────────────┘  └──────────┘  └───────────────────┘   │
└──────────────────────┬───────────────────────────────────┘
                       │ HTTPS (TLS 1.3)
┌──────────────────────▼───────────────────────────────────┐
│              POINT D'ENTRÉE (Nginx)                      │
│  • Terminaison TLS    • Headers sécurité    • CORS       │
│  • Reverse proxy vers backend                            │
└──────────────────────┬───────────────────────────────────┘
                       │ HTTP interne
┌──────────────────────▼───────────────────────────────────┐
│              BACKEND (FastAPI)                            │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐ │
│  │Validation│  │ Paillier │  │  Audit   │  │  IAM     │ │
│  │ Pydantic │  │ Agrégat° │  │Hash-Chain│  │ Vérif.   │ │
│  └──────────┘  └──────────┘  └──────────┘  └──────────┘ │
└──┬──────────────────────────────────────────────┬────────┘
   │                                              │
┌──▼──────────────────────┐  ┌────────────────────▼────────┐
│  PostgreSQL (4 silos)   │  │  MinIO (S3)                 │
│  ┌────────┐ ┌────────┐  │  │  Photos chiffrées           │
│  │identity│ │ votes  │  │  │  AES-256-GCM                │
│  │(hachés)│ │(Pailli)│  │  └─────────────────────────────┘
│  ├────────┤ ├────────┤  │
│  │ audit  │ │ photos │  │
│  │(chain) │ │(méta)  │  │
│  └────────┘ └────────┘  │
└─────────────────────────┘
"""
add_code_block(architecture_desc.strip())

doc.add_heading("2.3 Trois réseaux logiques", level=2)
add_table(
    ["Réseau", "Rôle", "Services concernés"],
    [
        ["Control Plane", "Monitoring, CI/CD, health checks", "GitHub Actions, scripts audit"],
        ["Application Plane", "Flux métier, API, logique applicative", "Nginx, FastAPI, modules crypto"],
        ["Data Plane", "Stockage données sensibles", "PostgreSQL (4 silos), MinIO"],
    ],
    col_widths=[4, 5, 6]
)

doc.add_heading("2.4 Justifications architecturales", level=2)

justifications = [
    ("Chiffrement au repos", "Les données stockées contiennent des informations personnelles protégées par le RGPD. Sans chiffrement au repos, un accès non autorisé au disque ou à une sauvegarde exposerait toutes les données en clair. Les votes sont chiffrés par Paillier dans les colonnes, les photos sont chiffrées AES-GCM dans MinIO."),
    ("Chiffrement homomorphe", "Le chiffrement homomorphe additif (Paillier) permet de calculer la somme des votes sans jamais déchiffrer les réponses individuelles. Seule l'institution habilitée, détenant la clé privée, peut déchiffrer le total agrégé, garantissant la confidentialité absolue du vote."),
    ("Séparation en silos", "La séparation en 4 schémas PostgreSQL distincts (identity, votes, audit, photos) applique le principe de séparation des responsabilités et limite l'impact d'une compromission : un accès au silo votes ne permet pas de corréler avec les identités."),
    ("Point d'entrée unique", "Le reverse proxy Nginx centralise la terminaison TLS, l'application des headers de sécurité, et le routage. Il constitue le seul point d'entrée, réduisant la surface d'attaque conformément au principe de défense en profondeur."),
    ("Backend en réseau isolé", "L'isolation réseau (networks Docker ou VPC cloud) empêche tout accès direct aux services internes depuis l'extérieur. Seul le reverse proxy est exposé publiquement."),
    ("Protection bout en bout", "Le vote est chiffré côté client (Paillier BigInt JS), la photo côté client (AES-256-GCM Web Crypto API), la transmission via HTTPS (TLS 1.3), le stockage dans des silos séparés. L'agrégation s'effectue sur les ciphertexts sans déchiffrement."),
]

for title, desc in justifications:
    p = doc.add_paragraph()
    run = p.add_run(f"{title} : ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 3. MODÈLE UML
# ═══════════════════════════════════════════════════

doc.add_heading("3. Modèle UML", level=1)

doc.add_heading("3.1 Diagramme de classes", level=2)
doc.add_paragraph(
    "Le système manipule quatre entités principales, réparties dans des schémas PostgreSQL distincts, "
    "ainsi que deux classes cryptographiques (Paillier) :"
)

add_table(
    ["Classe", "Schéma", "Attributs principaux"],
    [
        ["Participant", "identity", "id, participant_hash (SHA-256), agent_hash, created_at"],
        ["Vote", "votes", "id, question_id, participant_hash, ciphertext (Paillier), key_id, created_at"],
        ["AuditLog", "audit", "id, event_type, payload_hash, prev_hash (→ chaîne), created_at"],
        ["Photo", "photos", "id, object_name, nonce, tag, alg, content_type, size_bytes, key_id, created_at"],
        ["PaillierPublicKey", "—", "n (BigInt), g (BigInt) — méthodes : encrypt(m), add(c1,c2)"],
        ["PaillierPrivateKey", "—", "lam, mu, n (BigInt) — méthodes : decrypt(c)"],
    ],
    col_widths=[3.5, 2, 10]
)

doc.add_heading("3.2 Relations", level=2)
relations = [
    "Vote → Participant : via participant_hash (référence indirecte par hash)",
    "Vote → AuditLog : chaque soumission de vote crée une entrée d'audit",
    "Photo → AuditLog : chaque upload de photo crée une entrée d'audit",
    "AuditLog → AuditLog : chaîne via prev_hash (hash de l'entrée précédente)",
    "Vote → PaillierPublicKey : le ciphertext est chiffré avec la clé publique",
]
for r in relations:
    add_bullet(r)

doc.add_heading("3.3 Diagramme de séquence — Collecte d'un vote", level=2)
doc.add_paragraph(
    "Le workflow de collecte suit 6 étapes séquentielles, du scan QR à la transmission chiffrée :"
)

sequence = """
Agent Terrain (PWA)          Backend FastAPI          PostgreSQL          MinIO
       │                          │                      │                │
       │── Scanner QR Question ──►│                      │                │
       │◄── question_id ─────────│                      │                │
       │                          │                      │                │
       │── Scanner QR Participant►│                      │                │
       │◄── participant_id ──────│                      │                │
       │                          │                      │                │
       │── Vote (0 ou 1) ────────│                      │                │
       │── Chiffrer (Paillier) ──│                      │                │
       │                          │                      │                │
       │── Capturer photo ───────│                      │                │
       │── Chiffrer (AES-GCM) ──│                      │                │
       │                          │                      │                │
       │── POST /votes/send ─────►│                      │                │
       │                          │── INSERT vote ──────►│                │
       │                          │── INSERT identity ──►│                │
       │                          │── INSERT audit ─────►│                │
       │◄── {vote_id, stored} ───│                      │                │
       │                          │                      │                │
       │── POST /uploads/photo ──►│                      │                │
       │                          │── PUT objet ─────────────────────────►│
       │                          │── INSERT photo meta─►│                │
       │                          │── INSERT audit ─────►│                │
       │◄── {photo_id, stored} ──│                      │                │
"""
add_code_block(sequence.strip())

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 4. CHOIX TECHNOLOGIQUES
# ═══════════════════════════════════════════════════

doc.add_heading("4. Choix technologiques", level=1)

doc.add_heading("4.1 Application mobile — PWA", level=2)
add_table(
    ["Critère", "Choix : PWA (HTML/JS)", "Alternative : Flutter"],
    [
        ["Portabilité", "✅ Tous navigateurs (iOS/Android)", "✅ iOS + Android natif"],
        ["Crypto côté client", "✅ Web Crypto API + BigInt natif", "⚠️ Nécessite packages Dart"],
        ["Installation", "✅ Aucune (URL directe)", "❌ Stores obligatoires"],
        ["Caméra / QR", "✅ html5-qrcode (MediaDevices)", "✅ Plugins natifs"],
        ["Mode offline", "✅ localStorage + sync auto", "✅ Hive / SQLite"],
    ],
    col_widths=[3.5, 5.5, 5.5]
)
doc.add_paragraph(
    "Justification : La PWA offre un déploiement immédiat sans store, un accès natif à la caméra "
    "et aux API cryptographiques du navigateur. Pour un prototype, c'est le choix le plus pragmatique."
)

doc.add_heading("4.2 Backend — FastAPI (Python)", level=2)
add_table(
    ["Critère", "Choix : FastAPI", "Alternative : Spring Boot"],
    [
        ["Performance async", "✅ Natif (uvicorn/ASGI)", "✅ Threads JVM"],
        ["Validation entrées", "✅ Pydantic automatique", "✅ Bean Validation"],
        ["Doc API", "✅ OpenAPI/Swagger auto", "✅ SpringDoc"],
        ["Bibliothèques crypto", "✅ cryptography (mature)", "✅ Bouncy Castle"],
        ["Rapidité développement", "✅ Très rapide", "⚠️ Plus verbeux"],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

doc.add_heading("4.3 Base de données — PostgreSQL", level=2)
add_table(
    ["Critère", "Choix : PostgreSQL 15", "Alternative : MongoDB"],
    [
        ["Multi-schémas", "✅ Natif → 4 silos", "❌ Pas de schémas logiques"],
        ["Chiffrement", "✅ pgcrypto + TDE possible", "⚠️ Encryption at rest uniquement"],
        ["RBAC granulaire", "✅ Rôles natifs par schéma", "⚠️ Plus limité"],
        ["Transactions ACID", "✅ Complet", "⚠️ Eventual consistency"],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

doc.add_heading("4.4 Communication sécurisée", level=2)
add_table(
    ["Élément", "Choix", "Justification"],
    [
        ["Protocole", "TLS 1.3", "Standard actuel, forward secrecy, performance"],
        ["Certificats", "X.509", "Authentification mutuelle terminaux"],
        ["Reverse proxy", "Nginx", "Terminaison TLS, headers sécurité, performance"],
    ],
    col_widths=[3, 3, 9]
)

doc.add_heading("4.5 Environnement cloud", level=2)
doc.add_paragraph(
    "Docker Compose (cloud privé / local) est utilisé pour le prototypage. "
    "L'architecture est conçue pour être transposable vers AWS (ECS + RDS + S3), "
    "GCP (Cloud Run + Cloud SQL + GCS), ou Azure (ACI + Azure DB + Blob Storage)."
)
doc.add_paragraph(
    "L'isolation réseau (VPC ou networks Docker) est obligatoire car les bases de données "
    "et le stockage objet ne doivent jamais être accessibles depuis Internet. "
    "Seul le point d'entrée (Nginx / API Gateway) possède une adresse publique."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 5. CONCEPTION CRYPTOGRAPHIQUE
# ═══════════════════════════════════════════════════

doc.add_heading("5. Conception cryptographique", level=1)

doc.add_heading("5.1 Chiffrement homomorphe — Paillier", level=2)
doc.add_paragraph(
    "Le schéma de Paillier est un cryptosystème homomorphe additif qui permet "
    "d'additionner des messages chiffrés sans les déchiffrer."
)

doc.add_heading("Principe mathématique", level=3)
formulas = [
    "Génération : p, q premiers → n = p·q, λ = lcm(p-1, q-1), g = n+1, μ = L(g^λ mod n²)⁻¹ mod n",
    "Chiffrement : E(m) = g^m · r^n mod n²  (r aléatoire, copremier avec n)",
    "Déchiffrement : D(c) = L(c^λ mod n²) · μ mod n  où L(x) = (x-1)/n",
    "Addition homomorphe : E(m₁) × E(m₂) mod n² = E(m₁ + m₂)",
]
for f in formulas:
    add_bullet(f)

doc.add_heading("Choix Paillier vs ElGamal", level=3)
add_table(
    ["Critère", "Paillier", "ElGamal additif"],
    [
        ["Type", "Nativement additif", "Modification nécessaire"],
        ["Opération", "Multiplication ciphertexts", "Multiplication ciphertexts"],
        ["Déchiffrement total", "Direct", "Nécessite table log discret"],
        ["Complexité", "O(1) pour agrégation", "O(√T) pour déchiffrement (baby-step)"],
        ["Adapté aux votes binaires", "✅ Idéal", "⚠️ Moins efficace"],
    ],
    col_widths=[4, 5, 5.5]
)

doc.add_heading("Flux de chiffrement", level=3)
flow = """Vote client (m ∈ {0,1})
    │
    ▼  Chiffrement Paillier (clé publique, côté client JS)
    E(m) = g^m · r^n mod n²
    │
    ▼  Transmission HTTPS → Backend
    │
    ▼  Stockage : ciphertext dans silo "votes"
    │
    ▼  Agrégation : ∏ E(mᵢ) mod n² = E(∑ mᵢ)
    │
    ▼  Déchiffrement (clé privée institution) → total votes "Oui"
"""
add_code_block(flow.strip())

add_note_box(
    "L'implémentation utilise des clés de 256 bits (pédagogique). "
    "En production, il est impératif d'utiliser des clés ≥ 2048 bits "
    "avec une bibliothèque auditée (python-phe, SEAL, HElib).",
    "Avertissement"
)

doc.add_heading("5.2 Chiffrement des photos — AES-256-GCM", level=2)
add_table(
    ["Propriété", "Valeur"],
    [
        ["Algorithme", "AES-256-GCM (AEAD)"],
        ["Longueur clé", "256 bits"],
        ["Nonce (IV)", "12 octets aléatoires (crypto.getRandomValues)"],
        ["Tag authenticité", "128 bits (16 octets)"],
        ["API côté client", "Web Crypto API (SubtleCrypto)"],
        ["API côté serveur", "Python cryptography (tests)"],
        ["Standard", "NIST SP 800-38D"],
    ],
    col_widths=[5, 10]
)
doc.add_paragraph(
    "Le mode GCM (Galois/Counter Mode) fournit simultanément le chiffrement et l'authentification (AEAD). "
    "Toute altération est détectée lors du déchiffrement par vérification du tag."
)

doc.add_heading("5.3 Hashage des identifiants", level=2)
doc.add_paragraph(
    "Les identifiants des participants et agents sont hachés avec SHA-256 avant stockage. "
    "Cela permet de vérifier l'identité lors d'audits sans stocker les identifiants en clair, "
    "conformément au RGPD (pseudonymisation)."
)

doc.add_heading("5.4 Propriétés de sécurité garanties", level=2)
add_table(
    ["Propriété", "Mécanisme"],
    [
        ["Confidentialité des votes", "Chiffrement Paillier (jamais déchiffré individuellement)"],
        ["Confidentialité des photos", "AES-256-GCM côté client (nonce aléatoire)"],
        ["Intégrité des données", "Hash chain audit (SHA-256, prev_hash)"],
        ["Intégrité des photos", "Tag GCM 128 bits (AEAD)"],
        ["Anonymat des participants", "Hashage SHA-256 des identifiants"],
        ["Non-corrélation", "Silos séparés (identité ≠ votes ≠ photos)"],
        ["Non-répudiation", "Audit log avec agent_id + timestamp"],
    ],
    col_widths=[5, 10]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 6. SÉCURITÉ
# ═══════════════════════════════════════════════════

doc.add_heading("6. Sécurité", level=1)

doc.add_heading("6.1 Headers HTTP de sécurité", level=2)
add_table(
    ["Header", "Valeur", "Protection"],
    [
        ["Strict-Transport-Security", "max-age=63072000", "Force HTTPS pendant 2 ans"],
        ["X-Content-Type-Options", "nosniff", "Empêche le MIME sniffing"],
        ["X-Frame-Options", "DENY", "Empêche le clickjacking"],
        ["X-XSS-Protection", "1; mode=block", "Protection XSS navigateur"],
        ["Referrer-Policy", "strict-origin-when-cross-origin", "Limite les informations Referer"],
        ["Cache-Control", "no-store", "Empêche le cache de données sensibles"],
    ],
    col_widths=[4.5, 5, 5.5]
)

doc.add_heading("6.2 CORS", level=2)
doc.add_paragraph(
    "Le middleware CORS restreint les origines autorisées au frontend uniquement "
    "(http://localhost:3000), ne permet que les méthodes GET et POST, et interdit les credentials cross-origin."
)

doc.add_heading("6.3 Validation des entrées", level=2)
doc.add_paragraph(
    "Le framework Pydantic (intégré à FastAPI) valide automatiquement chaque requête : "
    "types, longueurs min/max, formats. Toute entrée invalide retourne une erreur 422 avec le détail."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 7. APPLICATION MOBILE (PWA)
# ═══════════════════════════════════════════════════

doc.add_heading("7. Application mobile (PWA)", level=1)

doc.add_heading("7.1 Architecture front-end", level=2)
add_table(
    ["Fichier", "Rôle"],
    [
        ["index.html", "Page principale — workflow 5 étapes"],
        ["css/style.css", "Design mobile-first, thème sombre premium"],
        ["js/app.js", "Logique métier, navigation, mode offline"],
        ["js/paillier.js", "Chiffrement Paillier côté client (BigInt natif JS)"],
        ["js/crypto.js", "Chiffrement AES-256-GCM (Web Crypto API)"],
        ["js/qr-scanner.js", "Scanner QR via caméra (html5-qrcode)"],
        ["js/api.js", "Client HTTP vers le backend"],
        ["qr-generator.html", "Utilitaire : génération de QR codes de test"],
    ],
    col_widths=[4, 11]
)

doc.add_heading("7.2 Workflow agent terrain", level=2)
workflow_steps = [
    ("Étape 1 — Scanner la question", "L'agent scanne le QR code de la question ou saisit l'ID manuellement."),
    ("Étape 2 — Scanner le participant", "L'agent scanne le QR code du participant."),
    ("Étape 3 — Enregistrer le vote", "Le participant répond Oui (1) ou Non (0) via deux boutons visuels."),
    ("Étape 4 — Prendre la photo", "Capture via la caméra du smartphone (pouce levé/baissé)."),
    ("Étape 5 — Résumé et envoi chiffré", "Le vote est chiffré (Paillier), la photo est chiffrée (AES-256-GCM), puis les données sont transmises au backend."),
]
for title, desc in workflow_steps:
    add_bullet(desc, bold_prefix=title)

doc.add_heading("7.3 Mode offline", level=2)
doc.add_paragraph(
    "Lorsque la connexion est indisponible, les données chiffrées sont stockées dans localStorage. "
    "Dès le retour en ligne, la synchronisation automatique transmet les votes en attente au backend. "
    "Un badge affiche le nombre de votes en file d'attente."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 8. BACKEND API REST
# ═══════════════════════════════════════════════════

doc.add_heading("8. Backend — API REST", level=1)

add_table(
    ["Méthode", "Endpoint", "Description"],
    [
        ["GET", "/health", "Health check (API + connectivité DB)"],
        ["POST", "/votes/send", "Soumettre un vote chiffré Paillier"],
        ["GET", "/votes/aggregate", "Agrégation homomorphe par question"],
        ["POST", "/uploads/photo", "Upload photo chiffrée AES-GCM"],
        ["GET", "/audit/verify", "Vérifier intégrité hash-chain"],
        ["GET", "/audit/logs", "Consulter les entrées d'audit"],
        ["GET", "/iam/verify", "Vérifier permissions DB (moindre privilège)"],
        ["GET", "/crypto/pubkey", "Clé publique Paillier (démo)"],
        ["POST", "/crypto/encrypt", "Chiffrer un vote côté serveur (fallback)"],
        ["GET", "/demo", "Page HTML interactive de test"],
    ],
    col_widths=[2, 4, 9]
)

doc.add_paragraph(
    "La documentation interactive Swagger est auto-générée par FastAPI et accessible à /docs."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 9. BASE DE DONNÉES
# ═══════════════════════════════════════════════════

doc.add_heading("9. Base de données — Séparation en silos", level=1)

doc.add_paragraph(
    "Les données sont réparties dans 4 schémas PostgreSQL distincts, chacun avec un rôle précis :"
)

add_table(
    ["Schéma", "Table", "Contenu", "Données sensibles"],
    [
        ["identity", "participants", "Hachés SHA-256 (participant + agent)", "Pseudonymisées"],
        ["votes", "votes", "question_id, ciphertext Paillier, key_id", "Chiffrées (Paillier)"],
        ["audit", "logs", "event_type, payload_hash, prev_hash (chaîne)", "Intégrité (hash)"],
        ["photos", "photos", "object_name, nonce, tag, alg, key_id", "Métadonnées crypto"],
    ],
    col_widths=[2.5, 3, 5.5, 3.5]
)

doc.add_paragraph(
    "Les photos elles-mêmes (données binaires chiffrées AES-GCM) sont stockées dans MinIO, "
    "séparé du PostgreSQL — constituant un 5ème silo physique isolé."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 10. AUDIT ET HASH-CHAIN
# ═══════════════════════════════════════════════════

doc.add_heading("10. Audit et journalisation immuable", level=1)

doc.add_heading("10.1 Principe de la hash-chain", level=2)
doc.add_paragraph(
    "Chaque entrée d'audit contient un payload_hash (SHA-256 du contenu de l'opération) "
    "et un prev_hash qui référence le payload_hash de l'entrée précédente. "
    "Toute modification ou suppression d'une entrée brise la chaîne et est détectable."
)

chain_example = """Entrée #1: payload_hash=abc123, prev_hash=NULL       (première)
Entrée #2: payload_hash=def456, prev_hash=abc123     (liée à #1)
Entrée #3: payload_hash=ghi789, prev_hash=def456     (liée à #2)
...
Si #2 est altérée → #3.prev_hash ≠ nouveau hash de #2 → CORRUPTION DÉTECTÉE"""
add_code_block(chain_example)

doc.add_heading("10.2 Événements audités", level=2)
add_table(
    ["Événement", "Déclencheur", "Payload hashé"],
    [
        ["vote_received", "POST /votes/send", "question_id:ciphertext:key_id"],
        ["photo_uploaded", "POST /uploads/photo", "photo:object_name:size:key_id"],
    ],
    col_widths=[4, 4, 7]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 11. CI/CD
# ═══════════════════════════════════════════════════

doc.add_heading("11. CI/CD", level=1)

doc.add_heading("11.1 Pipeline GitHub Actions", level=2)
doc.add_paragraph(
    "Le fichier .github/workflows/ci.yml définit un pipeline en 4 jobs parallèles, "
    "déclenchés sur chaque push ou pull request :"
)

add_table(
    ["Job", "Outil", "Objectif"],
    [
        ["lint", "ruff", "Qualité du code + règles de sécurité (bandit)"],
        ["test", "pytest", "Tests unitaires et d'intégration (PostgreSQL + MinIO)"],
        ["dependency-scan", "pip-audit", "Détection de vulnérabilités dans les dépendances"],
        ["secret-scan", "detect-secrets", "Détection de secrets (clés, mots de passe) dans le code"],
    ],
    col_widths=[3.5, 3, 8.5]
)

doc.add_paragraph(
    "Le merge n'est autorisé que si les 4 jobs passent. Le pipeline utilise des services "
    "PostgreSQL et MinIO identiques à l'environnement Docker Compose."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 12. IAM
# ═══════════════════════════════════════════════════

doc.add_heading("12. IAM — Gestion des accès", level=1)

doc.add_heading("12.1 Rôle base de données app_user", level=2)
doc.add_paragraph(
    "Le rôle app_user est créé avec les restrictions maximales (moindre privilège) :"
)
add_code_block(
    "CREATE ROLE app_user LOGIN PASSWORD '...'\n"
    "    NOSUPERUSER NOCREATEDB NOCREATEROLE NOINHERIT;"
)

doc.add_heading("12.2 Permissions par schéma", level=2)
add_table(
    ["Schéma", "SELECT", "INSERT", "UPDATE", "DELETE", "DROP"],
    [
        ["identity", "✅", "✅", "❌", "❌", "❌"],
        ["votes", "✅", "✅", "❌", "❌", "❌"],
        ["audit", "✅", "✅", "❌", "❌", "❌"],
        ["photos", "✅", "✅", "❌", "❌", "❌"],
    ],
    col_widths=[3, 2, 2, 2, 2, 2]
)
doc.add_paragraph(
    "Aucun droit UPDATE, DELETE ou DROP n'est accordé, garantissant l'immutabilité des données."
)

doc.add_heading("12.3 Vérification programmatique", level=2)
doc.add_paragraph(
    "L'endpoint /iam/verify interroge les tables systèmes PostgreSQL (pg_roles, has_table_privilege) "
    "pour vérifier dynamiquement que le rôle n'a pas de privilèges excessifs."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 13. TESTS ET AUDIT
# ═══════════════════════════════════════════════════

doc.add_heading("13. Tests et scripts d'audit", level=1)

doc.add_heading("13.1 Tests automatisés (pytest)", level=2)
add_table(
    ["Fichier de test", "Couverture"],
    [
        ["test_health.py", "Endpoint health, connectivité DB"],
        ["test_votes_send.py", "Soumission vote, stockage ciphertext, audit log, hashage"],
        ["test_votes_aggregate.py", "Agrégation homomorphe Paillier, cas sans vote"],
        ["test_upload_photo.py", "Upload photo AES-GCM, stockage MinIO, métadonnées"],
        ["test_audit_chain.py", "Hash-chain valide, détection corruption"],
        ["test_db_models.py", "Insertion/lecture dans les silos séparés"],
        ["test_paillier.py", "Encrypt/decrypt round-trip, addition homomorphe"],
        ["test_iam.py", "Vérification moindre privilège du rôle DB"],
    ],
    col_widths=[5, 10]
)

doc.add_heading("13.2 Scripts d'audit de sécurité", level=2)
doc.add_paragraph(
    "6 scripts automatisés vérifient la conformité continue du système :"
)
add_table(
    ["Script", "Vérification"],
    [
        ["check_votes_encrypted.py", "Aucun vote stocké en clair (0 ou 1 brut)"],
        ["check_photos_encrypted.py", "Chaque photo possède nonce, tag, key_id, algorithme valide"],
        ["check_iam.py", "Rôle DB non-superuser, grants conformes"],
        ["check_hash_chain.py", "Chaîne d'audit intègre (prev_hash correct)"],
        ["check_network.py", "Seuls les ports attendus sont exposés"],
        ["run_all_audits.py", "Lance tous les audits, produit un rapport JSON"],
    ],
    col_widths=[5.5, 9.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 14. CONFORMITÉ ET BILAN
# ═══════════════════════════════════════════════════

doc.add_heading("14. Conformité et bilan", level=1)

doc.add_heading("14.1 Conformité RGPD et ISO 27001", level=2)
add_table(
    ["Exigence", "Implémentation"],
    [
        ["RGPD — Pseudonymisation", "Hashage SHA-256 des identifiants participants"],
        ["RGPD — Minimisation", "Seules les données nécessaires sont collectées"],
        ["RGPD — Chiffrement", "Paillier (votes), AES-GCM (photos), séparation silos"],
        ["ISO 27001 — Contrôle accès", "IAM moindre privilège, rôles DB restreints"],
        ["ISO 27001 — Audit trail", "Hash-chain immuable, journalisation complète"],
        ["ISO 27001 — Gestion clés", "key_id versionné, rotation prévue"],
        ["Disponibilité", "Mode offline PWA, health checks Docker, restart auto"],
        ["Non-répudiation", "Audit logs avec agent_id + timestamp"],
    ],
    col_widths=[5, 10]
)

doc.add_heading("14.2 Bilan des fonctionnalités", level=2)
add_table(
    ["Fonctionnalité", "Statut"],
    [
        ["Collecte réponses binaires", "✅ Implémenté"],
        ["Scan QR codes (question + participant)", "✅ Implémenté"],
        ["Chiffrement homomorphe Paillier", "✅ Implémenté (client + serveur)"],
        ["Agrégation sans déchiffrement", "✅ Implémenté"],
        ["Chiffrement photos AES-256-GCM côté client", "✅ Implémenté"],
        ["Stockage séparé en silos", "✅ Implémenté (4 schémas + MinIO)"],
        ["Hashage identifiants", "✅ Implémenté (SHA-256)"],
        ["Journalisation immuable (hash-chain)", "✅ Implémenté"],
        ["IAM moindre privilège", "✅ Implémenté"],
        ["CI/CD pipeline", "✅ Implémenté (4 jobs)"],
        ["Scripts d'audit automatisés", "✅ Implémenté (6 scripts)"],
        ["Mode offline + synchronisation", "✅ Implémenté"],
        ["Headers de sécurité HTTP", "✅ Implémenté"],
        ["Documentation API (Swagger)", "✅ Auto-générée"],
        ["Tests automatisés", "✅ 8 fichiers de tests"],
    ],
    col_widths=[7, 8]
)

doc.add_paragraph()
doc.add_paragraph()

# Signature
sig = doc.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sig.add_run("— Fin du rapport —")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ═══════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    '..', 'docs', 'Rapport_Conception_Votes_Securises.docx'
)
output_path = os.path.normpath(output_path)
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"✅ Rapport généré : {output_path}")
