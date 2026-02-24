import os
import sys

def create_docx(output_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()

    # Title
    title = document.add_heading('Rapport de Conception & Réalisation — Système Sécurisé de Collecte de Votes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    document.add_heading('Introduction', level=1)
    document.add_paragraph(
        "Ce document valide les phases de conception et de développement du projet de système de collecte "
        "et d'agrégation homomorphe de votes sécurisé sur le terrain. Il répond point par point aux exigences techniques "
        "et de sécurité du cahier des charges."
    )

    # 1. Architecture Globale
    document.add_heading('1. Architecture globale de la solution', level=1)

    document.add_paragraph("[Insérer Schéma d'Architecture Cloud ici]")
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.paragraphs[-1].runs[0].font.color.rgb = RGBColor(255, 0, 0)
    document.paragraphs[-1].runs[0].italic = True

    document.add_paragraph("Conformément aux attentes, notre architecture est découpée en couches réseau isolées :")
    document.add_paragraph("Application Mobile (PWA) :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Responsable de la lecture des QR codes, de la capture photo, et du chiffrement local (Paillier pour le vote, AES-GCM pour la photo) avant transmission HTTPS.")
    
    document.add_paragraph("Accès Réseau (Reverse Proxy / Gateway) :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Le Nginx agit comme point d'entrée unique, assurant la terminaison TLS 1.3 et le routage vers le backend sécurisé.")
    
    document.add_paragraph("Backend (FastAPI) :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Serveur applicatif REST effectuant la validation stricte des requêtes (Pydantic), l'agrégation homomorphe, et l'insertion dans la journalisation immuable (Hash chain).")
    
    document.add_paragraph("Données (PostgreSQL & MinIO) :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" La base de données est séparée en silos (schémas 'identity', 'votes', 'audit'). Le stockage objet (MinIO) maintient les photos chiffrées strictement isolées (Data Plane).")

    # 2. Justification de l'Architecture (Questions du prof)
    document.add_heading("1.1 Justifications des choix d'architecture (Théorie)", level=2)

    p1 = document.add_paragraph("Pourquoi le chiffrement au repos est indispensable : ", style='List Bullet')
    p1.runs[0].bold = True
    p1.add_run("Les données collectées (photos, votes) relèvent de la vie privée (RGPD). Même si les accès réseau et applicatifs sont restreints, le chiffrement au repos garantit que si le support de stockage physique est volé ou si la base de données est dumpée, les données restent illisibles.")

    p2 = document.add_paragraph("Pourquoi le chiffrement homomorphe est utilisé pour les votes : ", style='List Bullet')
    p2.runs[0].bold = True
    p2.add_run("Il permet de sommer mathématiquement les réponses binaires chiffrées sans jamais déchiffrer le bulletin individuel. Ainsi, ni la base de données, ni l'administrateur système ne connaissent le vote d'un individu. Seule la personne possédant la clé privée finale obtient le résultat global, garantissant le secret total du scrutin.")

    p3 = document.add_paragraph("Pourquoi les données doivent être séparées en plusieurs silos : ", style='List Bullet')
    p3.runs[0].bold = True
    p3.add_run("C'est le principe de limitation de surface d'attaque et de non-corrélation. Si un attaquant parvient à compromettre la table des 'identités', il n'aura pas accès aux 'votes', et inversement. De plus, isoler les photos gérées par un stockage objet limite les vecteurs d'attaque sur la base relationnelle.")

    p4 = document.add_paragraph("Pourquoi un point d'entrée unique (API Gateway) est nécessaire : ", style='List Bullet')
    p4.runs[0].bold = True
    p4.add_run("Il centralise l'application des politiques de sécurité (Terminaison TLS 1.3, Headers de sécurité, CORS) et évite d'exposer publiquement les différents microservices internes. C'est la première ligne de défense (DMZ).")

    p5 = document.add_paragraph("Pourquoi le backend doit se situer dans un réseau isolé : ", style='List Bullet')
    p5.runs[0].bold = True
    p5.add_run("Le backend a un accès direct aux bases de données. L'isoler dans un réseau interne privé (Application Plane) garantit qu'il ne recevra de requêtes que via le point d'entrée unique contrôlé, le rendant invisible depuis l'Internet public.")

    p6 = document.add_paragraph("Comment les flux de données sont protégés de bout en bout : ", style='List Bullet')
    p6.runs[0].bold = True
    p6.add_run("1. Vote chiffré sur le mobile (Paillier). 2. Transmission dans un tunnel TLS (HTTPS). 3. API backend reçoit la donnée déjà chiffrée. 4. Stockage en l'état (au repos) dans des silos. 5. Agrégation homomorphe serveur sur les ciphertexts.")


    # 3. Choix Technologiques
    document.add_heading('2. Choix techniques et technologiques', level=1)
    
    document.add_heading("2.1 Côté Mobile", level=2)
    document.add_paragraph("Choix retenu : Progressive Web App (HTML/JS/WebCrypto)", style='List Bullet').runs[0].bold = True
    document.add_paragraph("Alternative comparée : Flutter", style='List Bullet').runs[0].bold = True
    document.add_paragraph("Justification :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Bien que Flutter permette une compilation native, une PWA permet de réagir au besoin de portabilité instantanée sans passer par les Stores Apple/Google. Côté sécurité, la PWA utilise la Web Crypto API native du navigateur (hautement sécurisée et auditée dans V8/SpiderMonkey) pour l'AES-GCM et le BigInt pour Paillier, sans s'appuyer sur des plugins tiers Flutter potentiellement vulnérables.")

    document.add_heading("2.2 Chiffrement Local", level=2)
    document.add_paragraph("Photos (AES-256-GCM) :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" L'AES-GCM est un algorithme de chiffrement authentifié (AEAD). Contrairement au CBC, le GCM chiffre ET crée un tag d'intégrité (128 bits). Si l'image chiffrée est altérée dans la BDD ou interceptée furtivement, le tag invalide le fichier lors du déchiffrement.")
    document.add_paragraph("Votes (Paillier) :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Nous avons opté pour le chiffrement homomorphe additif de Paillier. Justification : Paillier est nativement additif (la multiplication des ciphertexts donne l'addition des clairs), ce qui est algorithmiquement plus simple et efficace pour sommer des millions de bulletins O/1 que le ElGamal additif qui imposerait un problème de log discret en fin de dépouillement.")

    document.add_heading("2.3 Backend et Base de Données", level=2)
    document.add_paragraph("Backend (Python avec FastAPI) :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Comparé à SpringBoot ou ExpressJS, FastAPI apporte par conception la validation stricte des données (Pydantic). Il profite également de la bibliothèque de cryptographie asymétrique native la plus mûre (cryptography) idéale pour implémenter de grands entiers ou l'arithmétique Paillier. C'est asynchrone, rapide, et sûr (type-hinted).")
    
    document.add_paragraph("BDD (PostgreSQL) :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Par rapport à MongoDB, Postgres possède un contrôle d'accès natif basé sur les rôles (RBAC) extrêmement granulaire, et surtout un système de SCHÉMAS (namespaces) qui est parfait pour créer véritablement nos silos de données isolés (`identity`, `votes`, `audit`).")

    document.add_heading("2.4 Audit et Journalisation", level=2)
    document.add_paragraph("Hash Chaining :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Le système d'audit est construit comme une blockchain interne. Chaque insertion calcule un SHA-256 comprenant l'action ET le hash de l'action précédente (`prev_hash`). Pour garantir l'intégrité, aucune donnée ne peut être altérée a posteriori sans briser mathématiquement toute la chaîne consécutive de hachage. Un endpoint `/audit/verify` automatise cette vérification.")


    # 4. Phase 2 : Développement Professionnel & CI/CD
    document.add_heading('3. Phase 2 : Industrialisation et Sécurité Opérationnelle', level=1)

    document.add_heading("3.1 Organisation des 3 réseaux logiques", level=2)
    document.add_paragraph("Via Docker Compose (réseaux de type 'bridge' ou VPC cloud), nous avons isolé :", style='Normal')
    document.add_paragraph("Le Control Plane", style='List Bullet').runs[0].bold = True
    document.add_paragraph("L'Application Plane (Nginx vers FastAPI api_net)", style='List Bullet').runs[0].bold = True
    document.add_paragraph("Le Data Plane (Seule l'API parle à Postgres et MinIO db_net)", style='List Bullet').runs[0].bold = True

    document.add_heading("3.2 CI/CD", level=2)
    document.add_paragraph(
        "Un pipeline GitHub Actions a été mis en place pour automatiser le contrôle qualité et la sécurité avant déploiement. "
        "Les jobs incluent : 1. Linting strict (Ruff), 2. Tests (pytest avec postgres), 3. Scan de dépendances (pip-audit), "
        "4. Détection de fuites de secrets (detect-secrets). Toute PR contenant des vulnérabilités bloque le déploiement."
    )

    document.add_heading("3.3 IAM et Moindre Privilège", level=2)
    document.add_paragraph(
        "Nous avons appliqué le principe du moindre privilège jusqu'au niveau de la base de données. "
        "L'API FastAPI se connecte via le rôle (app_user) défini avec les attributs `NOSUPERUSER NOCREATEDB`. "
        "Des commandes `GRANT` strictes sont appliquées, et surtout : l'API a l'interdiction (`REVOKE`) de faire un `DELETE` "
        "sur la table des votes pour empêcher le serveur applicatif lui-même de falsifier l'historique en cas de prise de contrôle."
    )

    document.add_heading("3.4 Fonctionnalités de la GUI (Mobile & Admin)", level=2)
    document.add_paragraph("Côté Agent Terrain :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Workflow guidé en 5 étapes. Scan unitaire des QR questions et participants pour éviter la saisie (qui reste optionnelle avec validation API temps réel). Vote sécurisé masqué, prise de photo, et mode file d'attente hors-ligne en cas de perte de réseau.")
    
    document.add_paragraph("[Insérer Screenshot GUI Agent Terrain (QR et Vote) ici]")
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.paragraphs[-1].runs[0].font.color.rgb = RGBColor(255, 0, 0)
    document.paragraphs[-1].runs[0].italic = True

    document.add_paragraph("Côté Administrateur :", style='List Bullet').runs[0].bold = True
    document.paragraphs[-1].add_run(" Dashboard centralisé. Générateur de QR codes cryptographiques pour assigner des identités non prévisibles aux participants (existant via hash-chain asynchrone pour anonymat maximal). L'onglet clé est celui des Résultats avec déclenchement de l'Agrégation Homomorphe directement sur les données illisibles de la DB.")
    
    document.add_paragraph("[Insérer Screenshot GUI Administrateur (Générateur QR & Agrégation Homompore) ici]")
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.paragraphs[-1].runs[0].font.color.rgb = RGBColor(255, 0, 0)
    document.paragraphs[-1].runs[0].italic = True


    document.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    out_path = r"c:\Users\guigu\Downloads\sujet_app-main\sujet_app-main\Sujet_cnam\docs\Documentation_Finale_V2.docx"
    create_docx(out_path)
